"""
doc_generator.py — Generate professional Word .docx documents from schema data.

Takes a Schema and validated data dict, produces a python-docx Document with:
  - Professional title block with RFQ identification
  - Numbered sections for each content area
  - Formatted tables for work items, required documents, etc.
  - Compound fields rendered as labeled sub-sections
  - Conditional sections included/excluded based on data
"""

from __future__ import annotations

import io
from datetime import date, datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from engine.schema_loader import FieldDef, Schema

# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1F, 0x4E, 0x79)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------


def _setup_styles(doc: Document) -> None:
    """Define custom styles for consistent formatting.

    Args:
        doc: The Document to configure styles on.
    """
    style = doc.styles

    # Title style
    title_style = style["Title"]
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = NAVY
    title_style.font.name = "Calibri"

    # Heading 1
    h1 = style["Heading 1"]
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.font.name = "Calibri"

    # Heading 2
    h2 = style["Heading 2"]
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = DARK_GRAY
    h2.font.name = "Calibri"

    # Normal body text
    normal = style["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = "Calibri"


# ---------------------------------------------------------------------------
# Section counter
# ---------------------------------------------------------------------------


class _SectionCounter:
    """Auto-incrementing section number tracker."""

    def __init__(self) -> None:
        self._count = 0

    def next(self) -> int:
        """Return the next section number."""
        self._count += 1
        return self._count


# ---------------------------------------------------------------------------
# Document building helpers
# ---------------------------------------------------------------------------


def _add_header(doc: Document, data: dict[str, Any]) -> None:
    """Add the title block with RFQ identification and issuer info.

    Args:
        doc: The Document to add to.
        data: The field data dict.
    """
    # RFQ Title
    title = doc.add_paragraph("REQUEST FOR QUOTATION", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rfq_title = data.get("rfq_title", "")
    if rfq_title:
        sub = doc.add_paragraph(rfq_title)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = NAVY

    doc.add_paragraph("")

    # RFQ number and dates
    info_lines = []
    if data.get("rfq_number"):
        info_lines.append(f"RFQ Number: {data['rfq_number']}")
    if data.get("rfq_issue_date"):
        info_lines.append(f"Issue Date: {_format_date(data['rfq_issue_date'])}")
    if data.get("rfq_due_date"):
        due_time = data.get("rfq_due_time", "")
        due_str = f"Due Date: {_format_date(data['rfq_due_date'])}"
        if due_time:
            due_str += f" {due_time}"
        info_lines.append(due_str)

    for line in info_lines:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(10.5)

    doc.add_paragraph("")

    # Issuer info
    if data.get("issuer_name"):
        p = doc.add_paragraph("ISSUED BY:", style="Heading 2")
        doc.add_paragraph(data["issuer_name"])
        if data.get("issuer_address"):
            doc.add_paragraph(data["issuer_address"])

        contact_parts = []
        if data.get("issuer_contact_name"):
            name = data["issuer_contact_name"]
            if data.get("issuer_contact_title"):
                name += f", {data['issuer_contact_title']}"
            contact_parts.append(f"Contact: {name}")
        if data.get("issuer_contact_email"):
            contact_parts.append(f"Email: {data['issuer_contact_email']}")
        if data.get("issuer_contact_phone"):
            contact_parts.append(f"Phone: {data['issuer_contact_phone']}")

        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

    # Separator
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("")


def _add_section(
    doc: Document,
    counter: _SectionCounter,
    heading: str,
    content: str | None = None,
) -> None:
    """Add a numbered section with heading and optional body text.

    Args:
        doc: The Document to add to.
        counter: The section counter.
        heading: Section heading text.
        content: Optional body content.
    """
    num = counter.next()
    doc.add_heading(f"{num}. {heading}", level=1)
    if content:
        doc.add_paragraph(content)


def _add_table(
    doc: Document,
    field: FieldDef,
    rows: list[dict],
) -> None:
    """Add a formatted table from a table-type field.

    Args:
        doc: The Document to add to.
        field: The table field definition.
        rows: List of row dicts.
    """
    columns = field.columns or []
    if not columns or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for col_idx, col in enumerate(columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = col["label"]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
        # Navy background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(
            qn("w:shd"),
            {qn("w:fill"): "1F4E79", qn("w:val"): "clear"},
        )
        shading.append(shading_elm)

    # Data rows
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, col in enumerate(columns):
            value = row_data.get(col["key"], "")
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = _format_value_for_doc(col.get("type", "text"), value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")


def _add_compound_section(
    doc: Document,
    field: FieldDef,
    data: dict,
) -> None:
    """Render a compound field's sub-fields as labeled paragraphs.

    Args:
        doc: The Document to add to.
        field: The compound field definition.
        data: The compound field's value dict.
    """
    if not field.sub_fields or not isinstance(data, dict):
        return

    for sf in field.sub_fields:
        value = data.get(sf.key)
        if value and str(value).strip():
            p = doc.add_paragraph()
            run_label = p.add_run(f"{sf.label}: ")
            run_label.bold = True
            run_label.font.size = Pt(10.5)
            run_value = p.add_run(str(value))
            run_value.font.size = Pt(10.5)


def _should_include_section(field: FieldDef, data: dict[str, Any]) -> bool:
    """Check if a conditional field should be included based on data.

    Args:
        field: The field to check.
        data: The full data dict.

    Returns:
        True if the section should be included.
    """
    if not field.conditional_on:
        return True
    dep_field = field.conditional_on["field"]
    dep_value = field.conditional_on["value"]
    return data.get(dep_field) == dep_value


def _format_value_for_doc(field_type: str, value: Any) -> str:
    """Format a value for display in the document.

    Args:
        field_type: The field type string.
        value: The value to format.

    Returns:
        Formatted string.
    """
    if value is None:
        return ""
    if field_type == "boolean":
        return "Yes" if value else "No"
    if field_type == "currency":
        try:
            return f"${float(value):,.2f}"
        except (TypeError, ValueError):
            return str(value)
    if field_type == "date":
        return _format_date(value)
    return str(value)


def _format_date(value: Any) -> str:
    """Format a date value for display.

    Args:
        value: A date string or date object.

    Returns:
        Formatted date string.
    """
    if isinstance(value, (date, datetime)):
        return value.strftime("%B %d, %Y")
    if isinstance(value, str) and len(value) >= 10:
        try:
            dt = datetime.strptime(value[:10], "%Y-%m-%d")
            return dt.strftime("%B %d, %Y")
        except ValueError:
            pass
    return str(value)


# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------


def generate_document(
    schema: Schema,
    data: dict[str, Any],
    template_id: str | None = None,
) -> Document:
    """Generate a professional Word document from schema and data.

    Args:
        schema: The schema definition.
        data: Validated field data dict.
        template_id: Optional template ID (reserved for future use).

    Returns:
        A python-docx Document object ready to save.
    """
    doc = Document()
    _setup_styles(doc)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title block
    _add_header(doc, data)

    counter = _SectionCounter()

    # 1. Project Description
    if data.get("project_description"):
        _add_section(doc, counter, "PROJECT DESCRIPTION", data["project_description"])

        details = []
        if data.get("project_location"):
            details.append(f"Location: {data['project_location']}")
        if data.get("work_category"):
            details.append(f"Category: {data['work_category']}")
        if data.get("estimated_duration"):
            details.append(f"Duration: {data['estimated_duration']}")
        if data.get("estimated_start_date"):
            details.append(f"Start Date: {_format_date(data['estimated_start_date'])}")

        for detail in details:
            doc.add_paragraph(detail)

    # 2. Scope of Work
    if data.get("scope_summary"):
        _add_section(doc, counter, "SCOPE OF WORK", data["scope_summary"])

    # 2.1 Work Items table
    work_items_field = schema.get_field("work_items")
    if work_items_field and data.get("work_items"):
        doc.add_heading("Work Items", level=2)
        _add_table(doc, work_items_field, data["work_items"])

    # 2.2 Technical Specifications
    if data.get("specifications"):
        doc.add_heading("Technical Specifications", level=2)
        doc.add_paragraph(data["specifications"])

    # 3. Submission Requirements
    if data.get("submission_method"):
        _add_section(doc, counter, "SUBMISSION REQUIREMENTS")
        doc.add_paragraph(f"Method: {data['submission_method']}")
        if data.get("submission_address"):
            doc.add_paragraph(f"Address: {data['submission_address']}")

    # 3.1 Required Documents table
    req_docs_field = schema.get_field("required_documents")
    if req_docs_field and data.get("required_documents"):
        doc.add_heading("Required Documents", level=2)
        _add_table(doc, req_docs_field, data["required_documents"])

    # 4. Terms & Conditions
    terms_items = []
    if data.get("payment_terms"):
        terms_items.append(f"Payment: {data['payment_terms']}")
    if data.get("insurance_requirements"):
        terms_items.append(f"Insurance: {data['insurance_requirements']}")
    if "prevailing_wage" in data:
        pw = "Yes" if data["prevailing_wage"] else "No"
        terms_items.append(f"Prevailing Wage: {pw}")
    if "bonding_required" in data:
        br = "Yes" if data["bonding_required"] else "No"
        bond_str = f"Bond Required: {br}"
        if data.get("bonding_required") and data.get("bonding_amount"):
            bond_str += f" ({data['bonding_amount']})"
        terms_items.append(bond_str)

    if terms_items:
        _add_section(doc, counter, "TERMS & CONDITIONS")
        for item in terms_items:
            doc.add_paragraph(item)

    # 5. Pre-Bid Conference (if applicable)
    if data.get("prebid_conference"):
        _add_section(doc, counter, "PRE-BID CONFERENCE")
        if "prebid_mandatory" in data:
            mandatory = "Yes" if data["prebid_mandatory"] else "No"
            doc.add_paragraph(f"Mandatory: {mandatory}")
        if data.get("prebid_date"):
            doc.add_paragraph(f"Date/Time: {data['prebid_date']}")
        if data.get("prebid_location"):
            doc.add_paragraph(f"Location: {data['prebid_location']}")

    # 6. Evaluation Criteria (if data provided)
    eval_field = schema.get_field("evaluation_criteria")
    if eval_field and data.get("evaluation_criteria"):
        _add_section(doc, counter, "EVALUATION CRITERIA")
        _add_table(doc, eval_field, data["evaluation_criteria"])

    # 7. Additional Provisions (if applicable)
    has_additional = any(
        [
            data.get("safety_requirements"),
            data.get("environmental_requirements"),
            data.get("liquidated_damages"),
            data.get("retainage"),
        ]
    )

    if has_additional:
        _add_section(doc, counter, "ADDITIONAL PROVISIONS")

        # Safety requirements (compound)
        safety_field = schema.get_field("safety_requirements")
        if safety_field and data.get("safety_requirements"):
            doc.add_heading("Safety Requirements", level=2)
            _add_compound_section(doc, safety_field, data["safety_requirements"])

        if data.get("environmental_requirements"):
            doc.add_heading("Environmental Requirements", level=2)
            doc.add_paragraph(data["environmental_requirements"])

        if data.get("liquidated_damages"):
            doc.add_paragraph(f"Liquidated Damages: {data['liquidated_damages']}")

        if data.get("retainage"):
            doc.add_paragraph(f"Retainage: {data['retainage']}")

    # Flexible fields
    flex = data.get("_flexible_fields")
    if flex and isinstance(flex, list) and any(entry.get("field_value") for entry in flex):
        _add_section(doc, counter, "ADDITIONAL INFORMATION")
        for entry in flex:
            label = entry.get("field_label", "")
            value = entry.get("field_value", "")
            if value:
                p = doc.add_paragraph()
                run_label = p.add_run(f"{label}: ")
                run_label.bold = True
                p.add_run(str(value))

    return doc


# ---------------------------------------------------------------------------
# CLI testing
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    from engine.schema_loader import load_schema

    schema = load_schema("schemas/rfq_electric_utility.yaml")
    sample_data = {
        "issuer_name": "Test Utility Co",
        "rfq_number": "RFQ-TEST-001",
        "rfq_title": "Test RFQ Document",
        "rfq_issue_date": "2026-03-01",
        "rfq_due_date": "2026-03-28",
        "project_description": "Test project description.",
        "scope_summary": "Test scope.",
    }
    doc = generate_document(schema, sample_data)
    buf = io.BytesIO()
    doc.save(buf)
    print(f"Generated document: {buf.tell()} bytes")
