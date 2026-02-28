"""Tests for engine/doc_generator.py."""
from __future__ import annotations

import io

from docx import Document

from engine.doc_generator import generate_document
from engine.schema_loader import Schema


def _get_all_text(doc: Document) -> str:
    """Extract all text from a Document for assertion checks."""
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def test_generate_returns_document(rfq_schema: Schema, sample_data: dict) -> None:
    """generate_document returns a document with paragraphs."""
    doc = generate_document(rfq_schema, sample_data)
    assert hasattr(doc, "paragraphs")
    assert hasattr(doc, "tables")
    assert hasattr(doc, "save")
    assert len(doc.paragraphs) > 0


def test_document_contains_rfq_title(rfq_schema: Schema, sample_data: dict) -> None:
    """rfq_title appears in the document text."""
    doc = generate_document(rfq_schema, sample_data)
    text = _get_all_text(doc)
    assert "Distribution Line Reconstruction - Hwy 65 Corridor" in text


def test_document_contains_issuer(rfq_schema: Schema, sample_data: dict) -> None:
    """issuer_name appears in the document."""
    doc = generate_document(rfq_schema, sample_data)
    text = _get_all_text(doc)
    assert "Ozark Electric Cooperative" in text


def test_document_contains_work_items_table(rfq_schema: Schema, sample_data: dict) -> None:
    """Document has at least one table with work item headers."""
    doc = generate_document(rfq_schema, sample_data)
    assert len(doc.tables) >= 1

    # Find table with "Item #" header
    found = False
    for table in doc.tables:
        header_texts = [cell.text for cell in table.rows[0].cells]
        if "Item #" in header_texts:
            found = True
            break
    assert found


def test_document_contains_required_docs_table(rfq_schema: Schema, sample_data: dict) -> None:
    """Document has a table with 'Document' header (required documents)."""
    doc = generate_document(rfq_schema, sample_data)
    found = False
    for table in doc.tables:
        header_texts = [cell.text for cell in table.rows[0].cells]
        if "Document" in header_texts:
            found = True
            break
    assert found


def test_conditional_section_included(rfq_schema: Schema, sample_data: dict) -> None:
    """When bonding_required=True, bonding_amount appears in document."""
    data = {**sample_data, "bonding_required": True, "bonding_amount": "100% of contract value"}
    doc = generate_document(rfq_schema, data)
    text = _get_all_text(doc)
    assert "100% of contract value" in text


def test_conditional_section_excluded(rfq_schema: Schema, sample_data: dict) -> None:
    """When prebid_conference=False, prebid section is absent."""
    data = {**sample_data, "prebid_conference": False}
    doc = generate_document(rfq_schema, data)
    text = _get_all_text(doc)
    assert "PRE-BID CONFERENCE" not in text


def test_compound_section_rendered(rfq_schema: Schema, sample_data: dict) -> None:
    """Safety sub-fields appear as labeled content."""
    doc = generate_document(rfq_schema, sample_data)
    text = _get_all_text(doc)
    assert "General Safety Requirements:" in text
    assert "OSHA" in text


def test_optional_sections_only_when_data(rfq_schema: Schema, sample_data: dict) -> None:
    """Evaluation criteria only appears if data is provided."""
    # Without evaluation data
    doc_no_eval = generate_document(rfq_schema, sample_data)
    text_no = _get_all_text(doc_no_eval)
    assert "EVALUATION CRITERIA" not in text_no

    # With evaluation data
    data_with_eval = {
        **sample_data,
        "evaluation_criteria": [
            {"criterion": "Price", "weight": 40},
            {"criterion": "Experience", "weight": 60},
        ],
    }
    doc_with_eval = generate_document(rfq_schema, data_with_eval)
    text_yes = _get_all_text(doc_with_eval)
    assert "EVALUATION CRITERIA" in text_yes


def test_document_saveable(rfq_schema: Schema, sample_data: dict) -> None:
    """doc.save(BytesIO()) does not raise."""
    doc = generate_document(rfq_schema, sample_data)
    buf = io.BytesIO()
    doc.save(buf)
    assert buf.tell() > 0
